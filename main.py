import re
import sys
import time

from PyQt5 import QtWidgets, QtCore
from PyQt5.QtCore import QThread, pyqtSignal, pyqtSlot
from PyQt5.QtWidgets import QMainWindow, QApplication, QFrame, QFileDialog, QGraphicsDropShadowEffect
from ux import main_ux, file_frame, chostata_item_ux, file_item_ux, yondosh_word_ux, file_title_ux, file_btn
from pdfminer.high_level import extract_text
from docx import Document
from libs.Loading import LoadingScreen
from functools import partial


class Worker(QThread):
    set_file_signal = pyqtSignal(str)
    set_item_signal = pyqtSignal(str, str)
    set_end_signal = pyqtSignal(str)

    def __init__(self, parent):
        QThread.__init__(self, parent=parent)
        self.parent = parent

    def run(self):
        data_files = self.start_read_file()
        for data in data_files:
            self.set_file_signal.emit(data['file'])
            time.sleep(0.1)
            for key in sorted(data['words'].keys()):
                self.set_item_signal.emit(key, str(data['words'][key]))
                time.sleep(0.1)
        self.set_end_signal.emit("end")

    def stop(self):
        self.terminate()

    def word_count(self, str):
        counts = dict()
        words = str.split()

        for word in words:
            if word in counts:
                counts[word] += 1
            else:
                counts[word] = 1

        return counts

    def clean_text_only_words(self, text):
        # g‘ va o‘ harflarini vaqtincha almashtiramiz
        text = text.replace("g‘", "G_").replace("o‘", "O_")

        # Belgilarni tozalash — faqat o‘zbek va rus harflarini qoldiramiz
        # `'` va `’` bilan bir qatorda boshqa maxsus belgilar ham olib tashlanadi
        text = re.sub(r"[^a-zA-Zа-яА-ЯёЁўғқҳЎҒҚҲG_O_\s]", " ", text)

        # Ortiqcha bo'shliqlarni tozalaymiz
        text = re.sub(r'\s+', ' ', text).strip()

        # Yashirin belgilarni qayta almashtiramiz
        text = text.replace("G_", "g‘").replace("O_", "o‘")

        return text

    def start_read_file(self):
        file_data = []
        for file in self.parent.first_file:
            if file.endswith(".pdf"):
                text = extract_text(file)
                text = self.clean_text_only_words(text)
                text = text.lower()
                dt = {
                    "file": file.split("/")[-1],
                    "words": self.word_count(text)
                }
                file_data.append(dt)
            elif file.endswith(".docx"):
                text_doc = ""
                document = Document(file)
                for parag in document.paragraphs:
                    text_doc += "\n" + parag.text
                text_doc = text_doc.lower()
                text_doc = self.clean_text_only_words(text_doc)
                dt = {
                    "file": file.split("/")[-1],
                    "words": self.word_count(text_doc)
                }
                file_data.append(dt)

        return file_data

class WorkerStart(QThread):
    set_item_signal = pyqtSignal(str, str)
    set_end_signal = pyqtSignal(str)

    def __init__(self, parent):
        QThread.__init__(self, parent=parent)
        self.parent = parent

    def run(self):
        data_files = self.start_read_file()
        sort_keys = sorted(data_files.keys())
        for key in sort_keys:
            self.set_item_signal.emit(key, str(data_files[key]))
            time.sleep(0.1)
        self.set_end_signal.emit("end")

    def stop(self):
        self.terminate()

    def word_count(self, str):
        counts = dict()
        words = str.split()

        for word in words:
            if word in counts:
                counts[word] += 1
            else:
                counts[word] = 1

        return counts

    def clean_text_only_words(self, text):
        # g‘ va o‘ harflarini vaqtincha almashtiramiz
        text = text.replace("g‘", "G_").replace("o‘", "O_")

        # Belgilarni tozalash — faqat o‘zbek va rus harflarini qoldiramiz
        # `'` va `’` bilan bir qatorda boshqa maxsus belgilar ham olib tashlanadi
        text = re.sub(r"[^a-zA-Zа-яА-ЯёЁўғқҳЎҒҚҲG_O_\s]", " ", text)

        # Ortiqcha bo'shliqlarni tozalaymiz
        text = re.sub(r'\s+', ' ', text).strip()

        # Yashirin belgilarni qayta almashtiramiz
        text = text.replace("G_", "g‘").replace("O_", "o‘")

        return text

    def start_read_file(self):
        text_all = ""
        for file in self.parent.files:
            if file.endswith(".pdf"):
                text = extract_text(file).lower()
                text = self.clean_text_only_words(text)
                text_all += " " + text
            elif file.endswith(".docx"):
                text_doc = ""
                document = Document(file)
                for parag in document.paragraphs:
                    text_doc += "\n" + parag.text
                text_doc = self.clean_text_only_words(text_doc.lower())
                text_all += " " + text_doc

        words = self.word_count(text_all)
        return words




class WorkerSearch(QThread):
    set_file_chostata_signal = pyqtSignal(str, str, str)
    set_item_chostata_signal = pyqtSignal(str, str)
    set_yondosh_item_signal = pyqtSignal(str, str, str)
    finished = pyqtSignal(bool)

    def __init__(self, parent):
        QThread.__init__(self, parent=parent)
        self.parent = parent

    def run(self):
        self.search()
        self.finished.emit(True)

    def stop(self):
        self.terminate()

    def search(self):
        word = self.parent.word
        file_data = []
        total_text = ""
        for file in self.parent.files:
            if file.endswith(".pdf"):
                text = extract_text(file)
                text = text.lower()
                file = file.split("/")[-1]
                count = text.count(word)
                total_text += "\n" + text
                if count > 0:
                    self.set_file_chostata_signal.emit(file, str(count), word)
                time.sleep(0.1)
            elif file.endswith(".docx"):
                text_doc = ""
                document = Document(file)
                for parag in document.paragraphs:
                    text_doc += "\n" + parag.text
                text_doc = text_doc.lower()
                file = file.split("/")[-1]
                count = text_doc.count(word)
                total_text += "\n" + text_doc
                if count > 0:
                    self.set_file_chostata_signal.emit(file, str(count), word)
                time.sleep(0.1)
        self.checking_text(total_text, word)

    def checking_text(self, text: str, word: str):
        word = word.lower()
        text = text.lower()

        # Faqat so'zlar va tinish belgilarini ajratish
        sentences = re.split(r'(?<=[.!?])\s+', text)  # Jumlalarga bo'lish
        found_contexts = set()
        word_count = 0

        for sentence in sentences:
            words = sentence.split()
            for i, w in enumerate(words):
                if w == word:
                    word_count += 1

                    # Gap boshida va oxirida bo'lishiga qarab kontekstni aniqlash
                    if i == 0:
                        next_words = words[i + 1:i + 3]
                        context = f"{w} {' '.join(next_words)}"
                        prev_word = ""
                        next_word = " ".join(next_words)
                    elif i == len(words) - 1:
                        prev_words = words[max(0, i - 2):i]
                        context = f"{' '.join(prev_words)} {w}"
                        next_word = ""
                        prev_word = " ".join(prev_words)
                    else:
                        prev_word = words[i - 1]
                        next_word = words[i + 1] if i + 1 < len(words) else ""
                        context = f"{prev_word} {w} {next_word}"

                    # Takroriy kontekstlarni chiqarishni oldini olamiz
                    if context not in found_contexts:
                        found_contexts.add(context)
                        self.set_yondosh_item_signal.emit(prev_word, w, next_word)
                        time.sleep(0.1)

        # Umumiy so‘zlar sonini signalga uzatamiz
        if word_count > 0:
            self.set_item_chostata_signal.emit(word, str(word_count))



class FileFrame(QFrame, file_frame.Ui_Frame):
    def __init__(self, name):
        super(FileFrame, self).__init__()
        self.setupUi(self)
        self.file_name.setText(name)

        effect = QGraphicsDropShadowEffect()

        effect.setOffset(0, 0)

        effect.setBlurRadius(8)

        self.setGraphicsEffect(effect)


class ChostataItem(QFrame, chostata_item_ux.Ui_Frame):
    def __init__(self, word, count):
        super(ChostataItem, self).__init__()
        self.setupUi(self)
        self.label_word.setText(word)
        self.label_word_count.setText(str(count))
        effect = QGraphicsDropShadowEffect()

        effect.setOffset(0, 0)

        effect.setBlurRadius(5)

        self.setGraphicsEffect(effect)


class FileTitle(QFrame, file_title_ux.Ui_Frame):
    def __init__(self, file_name):
        super(FileTitle, self).__init__()
        self.setupUi(self)
        self.label_file.setText(file_name)
        effect = QGraphicsDropShadowEffect()

        effect.setOffset(0, 0)

        effect.setBlurRadius(5)

        self.setGraphicsEffect(effect)


class FileItem(QFrame, file_item_ux.Ui_Frame):
    def __init__(self, word, count, file):
        super(FileItem, self).__init__()
        self.setupUi(self)
        self.label_word.setText(word)
        self.label_word_count.setText(str(count))
        self.label_file.setText(file)
        effect = QGraphicsDropShadowEffect()

        effect.setOffset(0, 0)

        effect.setBlurRadius(5)

        self.setGraphicsEffect(effect)


class FileBtn(QFrame, file_btn.Ui_Frame):
    def __init__(self, file, parent):
        super(FileBtn, self).__init__(parent=parent)
        self.setupUi(self)
        self.pushButton.setText(file)
        effect = QGraphicsDropShadowEffect()

        effect.setOffset(0, 0)

        effect.setBlurRadius(5)

        self.setGraphicsEffect(effect)


class YondoshWord(QFrame, yondosh_word_ux.Ui_Frame):
    def __init__(self, word, word_before, word_after):
        super(YondoshWord, self).__init__()
        self.setupUi(self)
        self.label_word.setText(word)
        self.label_word_before.setText(word_before)
        self.label_word_after.setText(word_after)
        effect = QGraphicsDropShadowEffect()

        effect.setOffset(0, 0)

        effect.setBlurRadius(5)

        self.setGraphicsEffect(effect)


class Main(QMainWindow, main_ux.Ui_MainWindow):
    def __init__(self):
        super(Main, self).__init__()
        self.setupUi(self)
        try:
            self.loading = LoadingScreen(self)
            self.loading_timer = QtCore.QTimer()
            self.loading_timer.setInterval(500)
            self.loading_timer.timeout.connect(self.loadingAmimationActive)
            self.files = []
            self.word = ""
            self.first_file = [None]
            self.all_file = True
            self.effect = QGraphicsDropShadowEffect()
            self.stackedWidget.setCurrentWidget(self.page_file_upload)
            self.tabWidget.setCurrentWidget(self.tab)
            self.effect.setOffset(0, 0)
            self.effect.setBlurRadius(8)
            self.frame.setGraphicsEffect(self.effect)
            self.verticalLayout_side = QtWidgets.QVBoxLayout(self.scrollAreaWidgetContents_2)
            self.verticalLayout_right = QtWidgets.QVBoxLayout(self.scrollAreaWidgetContents)
            self.verticalLayout_chostata = QtWidgets.QVBoxLayout(self.scrollAreaWidgetContents_3)
            self.verticalLayout_file = QtWidgets.QVBoxLayout(self.scrollAreaWidgetContents_4)
            self.verticalLayout_yondosh = QtWidgets.QVBoxLayout(self.scrollAreaWidgetContents_5)
            self.upload_btn.clicked.connect(self.open_file)
            self.next_btn.clicked.connect(self.next_window_open)
            self.return_file_upload.clicked.connect(self.return_file_upload_window)
            self.search_btn.clicked.connect(self.search_word)
            self.file_render_worker = Worker(self)
            self.file_render_worker.set_file_signal.connect(self.set_card_file)
            self.file_render_worker.set_item_signal.connect(self.set_card_items)
            self.file_render_worker.set_end_signal.connect(self.set_card_end)
            self.start_item_render_worker = WorkerStart(self)
            self.start_item_render_worker.set_item_signal.connect(self.set_card_items)
            self.start_item_render_worker.set_end_signal.connect(self.set_card_end)
            self.search_render_worker = WorkerSearch(self)
            self.search_render_worker.set_file_chostata_signal.connect(self.set_file_items)
            self.search_render_worker.set_item_chostata_signal.connect(self.set_chostata_items)
            self.search_render_worker.set_yondosh_item_signal.connect(self.set_yondosh_word)
            self.search_render_worker.finished.connect(self.finished_search)

        except Exception as err:
            QtWidgets.QMessageBox.about(self, "Info", f'{err}')

    def loadingAmimationActive(self):
        self.loading.startAnimation()
        self.loading_timer.stop()

    def loadingAmimationStop(self):
        self.loading.stopAnimation()
        self.loading_timer.stop()

    @pyqtSlot(str, str)
    def set_card_items(self, word, count):
        v = ChostataItem(word, count)
        self.verticalLayout_chostata.addWidget(v)

    @pyqtSlot(str)
    def set_card_file(self, file):
        v = FileTitle(file)
        self.verticalLayout_chostata.addWidget(v)

    @pyqtSlot(str)
    def set_card_end(self, end):
        spacerItem = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_chostata.addItem(spacerItem)

    def search_word(self):
        word = self.lineEdit.text()
        word = word.lower()
        if word != "":
            self.file_render_worker.stop()
            self.search_render_worker.stop()
            self.start_item_render_worker.stop()
            self.remove_items(self.verticalLayout_chostata)
            self.remove_items(self.verticalLayout_yondosh)
            self.remove_items(self.verticalLayout_file)
            self.word = word
            self.search_render_worker.start()

    def open_file(self):
        path = QFileDialog.getOpenFileNames(self, 'Open a file', '',
                                            'Files (*.pdf;*.docx)')
        if path != ('', ''):
            self.remove_widgets()
            self.layout_set_widget(path[0])

    def next_window_open(self):
        if len(self.files) > 0:
            self.remove_files_side()
            self.set_files_side()
            self.stackedWidget.setCurrentWidget(self.page_main)
            # self.loading_timer.start()
            self.remove_items(self.verticalLayout_chostata)
            self.start_item_render_worker.start()

    def return_file_upload_window(self):
        self.stackedWidget.setCurrentWidget(self.page_file_upload)

    def click_file_side_btn(self, file):
        self.first_file[0] = file
        self.file_render_worker.stop()
        self.search_render_worker.stop()
        self.start_item_render_worker.stop()

        self.remove_items(self.verticalLayout_chostata)
        self.remove_items(self.verticalLayout_yondosh)
        self.remove_items(self.verticalLayout_file)
        self.file_render_worker.start()

    def set_files_side(self):

        for name in self.files:
            v = FileBtn(name.split("/")[-1], parent=self)
            v.pushButton.clicked.connect(partial(self.click_file_side_btn, name))
            self.verticalLayout_side.addWidget(v)
        if len(self.files) > 0:
            self.first_file[0] = self.files[0]
        spacerItem = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_side.addItem(spacerItem)

    def remove_widgets(self):
        self.files = []
        for i in reversed(range(self.verticalLayout_right.count())):
            w = self.verticalLayout_right.itemAt(i).widget()
            self.verticalLayout_right.removeWidget(w)

    def remove_files_side(self):
        for i in reversed(range(self.verticalLayout_side.count())):
            w = self.verticalLayout_side.itemAt(i).widget()
            self.verticalLayout_side.removeWidget(w)

    def layout_set_widget(self, names):
        self.files = names
        for name in names:
            v = FileFrame(name.split("/")[-1])
            self.verticalLayout_right.addWidget(v)
        spacerItem = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_right.addItem(spacerItem)

    @pyqtSlot(str, str)
    def set_chostata_items(self, word, count):
        v = ChostataItem(word, count)
        self.verticalLayout_chostata.addWidget(v)

    def remove_items(self, layout):
        if layout.count() > 0:
            for i in reversed(range(layout.count())):
                w = layout.itemAt(i).widget()
                if w is not None:
                    layout.removeWidget(w)

    @pyqtSlot(str, str, str)
    def set_file_items(self, file, count, word):
        v = FileItem(word, count, file)
        self.verticalLayout_file.addWidget(v)

    @pyqtSlot(str, str, str)
    def set_yondosh_word(self, word_before, word, word_after):
        v = YondoshWord(word, word_before, word_after)
        self.verticalLayout_yondosh.addWidget(v)

    @pyqtSlot(bool)
    def finished_search(self):
        spacerItem = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.verticalLayout_yondosh.addItem(spacerItem)
        self.verticalLayout_file.addItem(spacerItem)
        self.verticalLayout_chostata.addItem(spacerItem)


if __name__ == '__main__':
    app = QApplication(sys.argv)

    window = Main()
    window.show()

    sys.exit(app.exec_())
